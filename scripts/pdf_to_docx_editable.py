import re
import sys

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def overlaps(a, b):
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    return not (ax1 < bx0 or bx1 < ax0 or ay1 < by0 or by1 < ay0)


def clean_lines(text):
    lines = []
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if line:
            lines.append(line)

    result = []
    i = 0
    while i < len(lines):
        current = lines[i]
        nxt = lines[i + 1] if i + 1 < len(lines) else None
        if nxt == ".":
            fill = "_" * 22
            spacer = " " if current.endswith(":") else " : "
            result.append(f"{current}{spacer if not current.endswith(':') else ' '}{fill}")
            i += 2
            continue
        if current == ".":
            result.append("_" * 22)
        else:
            result.append(current)
        i += 1
    return result


def main():
    input_path, output_path = sys.argv[1:3]
    pdf = fitz.open(input_path)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    for page_index, page in enumerate(pdf):
        if page_index > 0:
            doc.add_page_break()

        tables = page.find_tables().tables if hasattr(page, "find_tables") else []
        table_bboxes = [tuple(t.bbox) for t in tables]
        table_map = {}

        for t in tables:
            data = t.extract()
            cleaned = []
            for row in data:
                if not row:
                    continue
                values = [re.sub(r"\s+", " ", (cell or "")).strip() for cell in row]
                if any(values):
                    cleaned.append(values)
            if cleaned:
                table_map[tuple(t.bbox)] = cleaned

        items = []
        for block in page.get_text("blocks"):
            bbox = tuple(block[:4])
            text = (block[4] or "").strip()
            if not text:
                continue
            if any(overlaps(bbox, tb) for tb in table_bboxes):
                continue
            items.append({"kind": "text", "y": bbox[1], "text": text})

        for bbox, data in table_map.items():
            items.append({"kind": "table", "y": bbox[1], "data": data})

        items.sort(key=lambda item: item["y"])

        for item in items:
            if item["kind"] == "text":
                lines = clean_lines(item["text"])
                joined = "    ".join(lines).strip()
                if not joined:
                    continue

                paragraph = doc.add_paragraph()
                if item["y"] < 150:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(joined)
                    run.bold = True
                    run.font.size = Pt(14 if item["y"] > 110 else 22)
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run(joined)
                    run.font.size = Pt(11)
                paragraph.paragraph_format.space_after = Pt(6)
            else:
                data = item["data"]
                cols = max(len(row) for row in data)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                for row in data:
                    cells = table.add_row().cells
                    for idx in range(cols):
                        cells[idx].text = row[idx] if idx < len(row) else ""

    doc.save(output_path)


if __name__ == "__main__":
    main()
