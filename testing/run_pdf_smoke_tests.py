from __future__ import annotations

import os
import json
from pathlib import Path
from typing import Any

import fitz
import requests
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
TESTING_DIR = ROOT / "testing"
OUTPUT_DIR = TESTING_DIR / "outputs"
RENDER_DIR = TESTING_DIR / "renders"
API_BASE = os.environ.get("PDF_SMOKE_API_BASE", "http://127.0.0.1:3000")
PASSWORD = "toolora123"


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    RENDER_DIR.mkdir(parents=True, exist_ok=True)


def post_file(endpoint: str, input_path: Path, output_path: Path, data: dict[str, Any] | None = None) -> requests.Response:
    with input_path.open("rb") as handle:
        response = requests.post(
            f"{API_BASE}{endpoint}",
            files={"file": (input_path.name, handle)},
            data=data or {},
            timeout=300,
        )
    output_path.write_bytes(response.content)
    return response


def post_merge(endpoint: str, input_paths: list[Path], output_path: Path) -> requests.Response:
    files = [("files", (path.name, path.open("rb"), "application/pdf")) for path in input_paths]
    try:
        response = requests.post(f"{API_BASE}{endpoint}", files=files, timeout=300)
        output_path.write_bytes(response.content)
        return response
    finally:
        for _, (_, handle, _) in files:
            handle.close()


def read_pdf(path: Path, password: str | None = None) -> dict[str, Any]:
    reader = PdfReader(str(path))
    encrypted = reader.is_encrypted
    if encrypted and password:
        reader.decrypt(password)
    text = reader.pages[0].extract_text() if reader.pages else ""
    return {
        "pages": len(reader.pages),
        "encrypted": encrypted,
        "text_preview": (text or "")[:400],
        "size": path.stat().st_size,
    }


def read_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    text = " | ".join(p.text for p in doc.paragraphs if p.text)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "text_preview": text[:400],
        "size": path.stat().st_size,
    }


def render_first_page(pdf_path: Path, image_path: Path) -> None:
    doc = fitz.open(str(pdf_path))
    pix = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    pix.save(str(image_path))


def write_markdown(results: dict[str, Any]) -> None:
    lines = [
        "# PDF Smoke Test Report",
        "",
        f"- API base: `{API_BASE}`",
        f"- Source DOCX: `{results['source_docx']}`",
        "",
        "## Results",
        "",
    ]

    for item in results["checks"]:
        status = "PASS" if item["ok"] else "FAIL"
        lines.append(f"### {item['name']} - {status}")
        lines.append("")
        lines.append(f"- Output: `{item['output']}`")
        if "details" in item:
            for key, value in item["details"].items():
                lines.append(f"- {key}: `{value}`")
        if item.get("note"):
            lines.append(f"- Note: {item['note']}")
        lines.append("")

    (TESTING_DIR / "TEST-REPORT.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    source_docx = TESTING_DIR / "Marksheet-example.docx"
    if not source_docx.exists():
        raise SystemExit(f"Missing source file: {source_docx}")

    health = requests.get(f"{API_BASE}/api/health", timeout=30)
    health.raise_for_status()

    checks: list[dict[str, Any]] = []

    converted_pdf = OUTPUT_DIR / "Marksheet-example.converted.pdf"
    response = post_file("/api/office/word-to-pdf", source_docx, converted_pdf)
    word_to_pdf_details = read_pdf(converted_pdf)
    render_first_page(converted_pdf, RENDER_DIR / "Marksheet-example.converted.page1.png")
    checks.append(
        {
            "name": "Word to PDF",
            "ok": response.ok and word_to_pdf_details["pages"] >= 1,
            "output": str(converted_pdf),
            "details": word_to_pdf_details,
        }
    )

    protected_pdf = OUTPUT_DIR / "Marksheet-example.protected.pdf"
    response = post_file("/api/pdf/protect", converted_pdf, protected_pdf, {"password": PASSWORD})
    protect_details = read_pdf(protected_pdf, PASSWORD)
    checks.append(
        {
            "name": "Protect PDF",
            "ok": response.ok and protect_details["encrypted"],
            "output": str(protected_pdf),
            "details": protect_details,
        }
    )

    unlocked_pdf = OUTPUT_DIR / "Marksheet-example.unlocked.pdf"
    response = post_file("/api/pdf/unlock", protected_pdf, unlocked_pdf, {"password": PASSWORD})
    unlock_details = read_pdf(unlocked_pdf)
    render_first_page(unlocked_pdf, RENDER_DIR / "Marksheet-example.unlocked.page1.png")
    checks.append(
        {
            "name": "Unlock PDF",
            "ok": response.ok and not unlock_details["encrypted"],
            "output": str(unlocked_pdf),
            "details": unlock_details,
        }
    )

    roundtrip_docx = OUTPUT_DIR / "Marksheet-example.roundtrip.docx"
    response = post_file("/api/office/pdf-to-word", converted_pdf, roundtrip_docx)
    roundtrip_docx_details = read_docx(roundtrip_docx)
    checks.append(
        {
            "name": "PDF to Word",
            "ok": response.ok and (
                roundtrip_docx_details["tables"] >= 1
                or roundtrip_docx_details["paragraphs"] >= 3
            ),
            "output": str(roundtrip_docx),
            "details": roundtrip_docx_details,
            "note": "Editable DOCX check: at least one table or multiple paragraphs should be present.",
        }
    )

    merged_pdf = OUTPUT_DIR / "merged-test.pdf"
    response = post_merge("/api/pdf/merge", [converted_pdf, unlocked_pdf], merged_pdf)
    merge_details = read_pdf(merged_pdf)
    checks.append(
        {
            "name": "Merge PDF",
            "ok": response.ok and merge_details["pages"] == 2,
            "output": str(merged_pdf),
            "details": merge_details,
        }
    )

    split_pdf = OUTPUT_DIR / "split-test.pdf"
    response = post_file("/api/pdf/split", converted_pdf, split_pdf)
    split_details = read_pdf(split_pdf)
    checks.append(
        {
            "name": "Split PDF",
            "ok": response.ok and split_details["pages"] == 1,
            "output": str(split_pdf),
            "details": split_details,
        }
    )

    rotated_pdf = OUTPUT_DIR / "rotated-test.pdf"
    response = post_file("/api/pdf/rotate", converted_pdf, rotated_pdf, {"rotation": "90"})
    rotate_details = read_pdf(rotated_pdf)
    checks.append(
        {
            "name": "Rotate PDF",
            "ok": response.ok and rotate_details["pages"] == 1,
            "output": str(rotated_pdf),
            "details": rotate_details,
        }
    )

    roundtrip_back_pdf = OUTPUT_DIR / "Marksheet-example.roundtrip.back-to-pdf.pdf"
    response = post_file("/api/office/word-to-pdf", roundtrip_docx, roundtrip_back_pdf)
    roundtrip_back_details = read_pdf(roundtrip_back_pdf)
    render_first_page(roundtrip_back_pdf, RENDER_DIR / "Marksheet-example.roundtrip.back-to-pdf.page1.png")
    checks.append(
        {
            "name": "Roundtrip Word to PDF",
            "ok": response.ok and roundtrip_back_details["pages"] >= 1,
            "output": str(roundtrip_back_pdf),
            "details": roundtrip_back_details,
        }
    )

    optional_office_samples = [
        (
            "PowerPoint to PDF",
            "/api/office/ppt-to-pdf",
            TESTING_DIR / "sample.pptx",
            OUTPUT_DIR / "sample.pptx.converted.pdf",
        ),
        (
            "Excel to PDF",
            "/api/office/excel-to-pdf",
            TESTING_DIR / "sample_input.xlsx",
            OUTPUT_DIR / "sample_input.xlsx.converted.pdf",
        ),
        (
            "HTML to PDF",
            "/api/office/html-to-pdf",
            TESTING_DIR / "sample.html",
            OUTPUT_DIR / "sample.html.converted.pdf",
        ),
    ]

    for name, endpoint, source_path, output_path in optional_office_samples:
        if not source_path.exists():
            checks.append(
                {
                    "name": name,
                    "ok": False,
                    "output": str(output_path),
                    "details": {"missing_source": str(source_path)},
                    "note": "Sample file not found in testing folder.",
                }
            )
            continue

        response = post_file(endpoint, source_path, output_path)
        details = read_pdf(output_path)
        checks.append(
            {
                "name": name,
                "ok": response.ok and details["pages"] >= 1,
                "output": str(output_path),
                "details": details,
            }
        )

    results = {
        "source_docx": str(source_docx),
        "health": health.json(),
        "checks": checks,
    }

    (TESTING_DIR / "test-results.json").write_text(json.dumps(results, indent=2), encoding="utf-8")
    write_markdown(results)
    print(json.dumps(results, indent=2))


if __name__ == "__main__":
    main()
